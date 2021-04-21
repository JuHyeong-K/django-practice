from docx import Document

document = Document()

document.add_heading('hi',0)
document.add_paragraph('heeeellllloo')

document.save('domo.docx')