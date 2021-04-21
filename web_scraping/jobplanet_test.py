import requests
import json
from bs4 import BeautifulSoup
from docx import Document

headers = {
    'authority': 'www.jobplanet.co.kr',
    'sec-ch-ua': '" Not A;Brand";v="99", "Chromium";v="90", "Google Chrome";v="90"',
    'accept': 'application/json, text/plain, */*',
    'x-newrelic-id': 'UgIBVlRRGwAHXFhWAQg=',
    'user_name': 'brain',
    'sec-ch-ua-mobile': '?0',
    'user-agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/90.0.4430.72 Safari/537.36',
    'sec-fetch-site': 'same-origin',
    'sec-fetch-mode': 'cors',
    'sec-fetch-dest': 'empty',
    'referer': 'https://www.jobplanet.co.kr/job/search?q=&page=1&occu_ids%5B%5D=11604&yoe=0%3B0',
    'accept-language': 'ko-KR,ko;q=0.9,en-US;q=0.8,en;q=0.7',
    'cookie': '__auc=b864817717732c43f5fb0cd237b; _ga=GA1.3.552040647.1611464458; _fbp=fb.2.1611464458521.1362224371; __gads=ID=f45e0f0d2bd39ede:T=1611470683:S=ALNI_MaNISJ4pRKQ8O1elwriY3YtAgp3Og; premiumInfoModalOn_190611=on; _jp_cookie_job_searching_modal_seen=true; __cfduid=d7ced2a65a41773fe34bf5dbf8cd474ad1618470596; hide_landing=true; _jp_visitor_token=9008e12b-3126-43d3-9379-c2f8fe9a5bc1; Jobplanet_remember_user_token=W1szNDQ1OTIxXSwiJDJhJDEwJHVoQ2kvNU5vUklwZjE4MFBCd3d1eWUiXQ%3D%3D--7b6e0560406815fa2bff852d828e3417ff9bdb15; _jp_cookie_subscribe_notice_email=true; _gid=GA1.3.588950140.1618830454; _jp_induction_status=true; _jp_visit_token=13711046-8d88-4936-bfc9-de18751c56f3; _jp_traffic_source=%7B%22utm_campaign%22%3Anull%2C%22utm_medium%22%3Anull%2C%22utm_source%22%3Anull%7D; _jp_visit_short_token=1618892201820-8332efb8-e176-4faa-876b-695cbb938f9c; __asc=7c99fbce178ed7ea5915b73d7a8; gnbTooltip_20201208=on; __cf_bm=4e1c49bb07407aaaf9d5f11ebcd8fafa4e1b2367-1618893145-1800-AVvthq4JJT131miBX4DPEnLhlscr4tv5RXdkyoc26J+2aydA+YunbzgYcTvhl0fv11U95+DtpP/ciAkO7EN8X5o=; _gat_UA-48990780-1=1; _intween_x2o_net_session=aEpKUHM2SXpnSVBlUk5XQmFBYU1NUzFETGY5S3hOVEhpTHJuRjgyZlkzbEJyVno5UzNFSy8yU0M1VFl1M01IcSt4RElZVDNFREVnMEhwcThqRzRpSDZuM2ZsSU9lUThrMWJUN25ZWVNFNFRtdHlNZFFiVTRBbnpvdU05NVJLTkllczFwVXRVeWwzTUUwd1dLMy84eUhvQnl6RDNqVnB2anh2cWJHbHJYUmdjdHROSGxKZVF1S0hFUnhuS3VHZmdzNHZrTU5PWDdMbG00YkNaclR6dVNGdFRJeDJSVXNQUEt2YUhwR3g1WUw0aENQeCtuZ096MUdUNHBBbTk1Mll6S2FxK202Z2dFYWlLZFNwZktFbU9idUxYNHhCMDB4dG9ObXBkTGIzaHpEdnBJRjFNZ2JLejFlWjE1WDBZN3VCUVZHUG9yckkwNk1sV2c4M2FUa1Z0eDA5WTRxUnp3TUhiemVDa2F6andEVzkwPS0tUGdlSy9aK2xFOEJVbHUxWGRXTzdKZz09--dbbf1ac01957e7a44f6b93623658ec99855a0a3a; wcs_bt=s_2f6e701a2a7e:1618893523',
    'if-none-match': 'W/"0f2ecff68067f2f4a17bebda2001bfa2"',
}
params = (
    ('q', ''),
    ('page', '1'),
    ('occu_ids[]', '11604'),
    ('yoe', '0;0'),
)

response = requests.get('https://www.jobplanet.co.kr/api/v3/job/search', headers=headers, params=params)
data = json.loads(response.text)
end_page = data['data']['display_pages'][4]['value']

document = Document()

for page in range(end_page):
    print('--------------------------------------')
    print(page)
    print('--------------------------------------')
    params = (
        ('q', ''),
        ('page', 1 + page),
        ('occu_ids[]', '11604'),
        ('yoe', '0;0'),
    )

    response = requests.get('https://www.jobplanet.co.kr/api/v3/job/search', headers=headers, params=params)
    data = json.loads(response.text)
    
    for i in range(len(data['data']['search_result']['jobs'])):
        print(data['data']['search_result']['jobs'][i]['company']['name'])
        print(data['data']['search_result']['jobs'][i]['jd']['title'])
        document.add_heading(data['data']['search_result']['jobs'][i]['company']['name'], 0)
        document.add_heading(data['data']['search_result']['jobs'][i]['jd']['title'], 1)
        try:
            print([i['name'] for i in data['data']['search_result']['jobs'][i]['jd']['required_skills']])
            required_skills = [i['name'] for i in data['data']['search_result']['jobs'][i]['jd']['required_skills']]
            document.add_paragraph(f'required_skills : {required_skills}')
        except:
            pass
        job_url = data['data']['search_result']['jobs'][i]['jd']['url']
        next_response = requests.get(f'https://www.jobplanet.co.kr/{job_url}')
        soup = BeautifulSoup(next_response.text, 'html.parser')

        main_task_info = soup.find('h3', string='주요업무')
        requirements_info = soup.find('h3', string='자격요건')
        prefer_info = soup.find('h3', string='우대사항')
        if main_task_info:
            print(main_task_info.find_next_sibling('div').get_text(separator='\n').strip())
            document.add_heading('주요업무', 2)
            a = main_task_info.find_next_sibling('div').get_text(separator='\n').strip()
            document.add_paragraph(f'{a}')
        if requirements_info:
            print(requirements_info.find_next_sibling('div').find('span').get_text(separator='\n').strip())
            document.add_heading('자격요건', 2)
            b = requirements_info.find_next_sibling('div').find('span').get_text(separator='\n').strip()
            document.add_paragraph(f'{b}')
        if prefer_info:
            print(prefer_info.find_next_sibling('div').get_text(separator='\n').strip())
            document.add_heading('우대사항', 2)
            c = prefer_info.find_next_sibling('div').get_text(separator='\n').strip()
            document.add_paragraph(f'{c}')
        document.add_page_break()

document.save('jobplanet_data.docx')

